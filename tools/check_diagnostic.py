# -*- coding: utf-8 -*-
"""Check suite for the diagnostic and adaptive system.

Everything here is arithmetic or a file-existence test. Nothing is a matter of
opinion, which is the point: the parts of this system that CAN be checked
mechanically should be, so that the parts that cannot get the attention.

    python3 tools/check_diagnostic.py            structure, marks, audio on disk
    python3 tools/check_diagnostic.py --probe    also re-probe every MP3 with ffprobe
"""
import os, re, sys, json, argparse, subprocess

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

from curriculum import all_lessons, teaching_lessons, load_papers, load_adaptive, load_profile
from curriculum.diagnostic import PAPERS, PAPER_AUDIO, DIAGNOSTIC_BLOCK
from curriculum.audio_diagnostic import DIAG_AUDIO, DIAG_FILES
from curriculum.rubrics import ALL_RUBRICS, BY_NAME

FAILS, WARNS, OKS = [], [], []
def ok(m):   OKS.append(m)
def warn(m): WARNS.append(m)
def fail(m): FAILS.append(m)


def check_sequence():
    ls = all_lessons()
    periods = [l.period for l in ls]
    if len(ls) != 94:
        fail(f"expected 94 sessions, found {len(ls)}")
    else:
        ok("94 sessions (92 taught + 2 diagnostic)")
    if periods != list(range(1, len(ls) + 1)):
        dupes = sorted({p for p in periods if periods.count(p) > 1})
        gaps = [p for p in range(1, len(ls) + 1) if p not in periods]
        fail(f"period sequence is not 1..{len(ls)} — duplicates {dupes}, missing {gaps}")
    else:
        ok(f"periods run 1..{len(ls)} with no gaps and no duplicates")
    if len(teaching_lessons()) != 92:
        fail(f"expected 92 taught sessions, found {len(teaching_lessons())}")
    else:
        ok("92 taught sessions unchanged")
    first = ls[:2]
    if [l.code for l in first] != ["D0L1", "D0L2"]:
        fail(f"the year does not open with the diagnostic: {[l.code for l in first]}")
    else:
        ok("the year opens with D0L1 and D0L2")
    for L in DIAGNOSTIC_BLOCK.lessons:
        mins = sum(s.minutes for s in L.procedure)
        if mins != 45:
            fail(f"{L.code}: procedure totals {mins} minutes, not 45")
    if not FAILS:
        ok("both diagnostic sessions are timed to exactly 45 minutes")


def check_marks():
    for P in load_papers():
        if abs(P.total - P.counted) > 1e-9:
            fail(f"Paper {P.code}: stated total {P.total} but items sum to {P.counted}")
        else:
            ok(f"Paper {P.code}: {P.total:g} marks, and the items sum to {P.counted:g}")
        for s in P.sections:
            if abs(s.marks - s.counted) > 1e-9:
                fail(f"Paper {P.code} section {s.code}: stated {s.marks} but items sum "
                     f"to {s.counted}")
        # written minutes must fit their period
        per = {}
        for s in P.sections:
            if s.strand in ("speaking", "pron"):
                continue
            per[s.period] = per.get(s.period, 0) + s.minutes
        for period, mins in per.items():
            if mins > 45:
                fail(f"Paper {P.code}: period {period} carries {mins} written minutes (max 45)")
        if not [f for f in FAILS if f"Paper {P.code}: period" in f]:
            ok(f"Paper {P.code}: written minutes fit their periods {dict(per)}")


def check_items():
    for P in load_papers():
        for s in P.sections:
            for t in s.tasks:
                if not t.items:
                    fail(f"{t.code}: no items")
                for it in t.items:
                    if not it.answer:
                        fail(f"{t.code}.{it.n}: no answer")
                    if not it.tests:
                        warn(f"{t.code}.{it.n}: no `tests` construct recorded")
                    parts = [x.strip() for x in re.split(r"[/→]", it.band)]
                    if any(x and x not in ("pre-A1", "A1", "A1+", "A2", "—") for x in parts):
                        warn(f"{t.code}.{it.n}: unrecognised band {it.band!r}")
                    if it.options and not any(it.answer.strip().startswith(o.strip()[0])
                                              for o in it.options):
                        warn(f"{t.code}.{it.n}: answer does not look like one of the options")
                if t.rubric and t.rubric not in BY_NAME:
                    fail(f"{t.code}: names rubric {t.rubric!r}, which does not exist")
                # a word box that promises N extra words must actually have N
                if t.wordbank:
                    m = re.search(r"(ONE|TWO|THREE|FOUR)\s+extra", t.instruction, re.I)
                    if m:
                        claimed = {"one": 1, "two": 2, "three": 3, "four": 4}[m.group(1).lower()]
                        actual = len(t.wordbank) - len(t.items)
                        if actual != claimed:
                            fail(f"{t.code}: instruction promises {claimed} extra word(s) but the "
                                 f"box has {len(t.wordbank)} words for {len(t.items)} gaps "
                                 f"({actual} extra)")
                    # Every answer must be gettable from the box. Compare on bare
                    # words, because a matching task labels its box "A. kitchen"
                    # while the key reads "A (kitchen)".
                    def _words(x):
                        return {w for w in re.findall(r"[a-z']+", x.lower()) if len(w) > 1}
                    box = set()
                    for w in t.wordbank:
                        box |= _words(w)
                    for it in t.items:
                        if it.options or not it.answer:
                            continue
                        got = _words(it.answer)
                        if got and not (got & box):
                            warn(f"{t.code}.{it.n}: answer {it.answer!r} shares no word with "
                                 f"the word box")
    if not [f for f in FAILS if "no answer" in f]:
        ok("every diagnostic item has an answer")
    if not [f for f in FAILS if "does not exist" in f]:
        ok("every rubric named by a task exists")
    if not [f for f in FAILS if "extra word" in f]:
        ok("every word box holds exactly as many spare words as its instruction promises")
    if not [w for w in WARNS if "word box" in w]:
        ok("every word-box answer is available in its own box")


# Papers A and C are placement instruments and must discriminate at the bottom, so
# at least 45% of their objective items sit at or below A1.  Paper B is a different
# instrument: it measures movement on content that has actually been taught in Units
# 1-6, so its floor is deliberately higher.  It still needs enough floor to re-band
# a student who is still Foundation in January, hence 30% rather than nothing.
FLOOR_TARGET = {"A": 0.45, "B": 0.30, "C": 0.45}


def check_calibration():
    """The floor must be weighted. That is the design claim; check it."""
    for P in load_papers():
        objective = [i for s in P.sections for t in s.tasks for i in t.items if not t.rubric]
        if not objective:
            continue
        low = sum(1 for i in objective if i.band in ("pre-A1", "A1"))
        share = low / len(objective)
        target = FLOOR_TARGET.get(P.code, 0.45)
        if share < target:
            fail(f"Paper {P.code}: only {share:.0%} of objective items are pre-A1/A1, "
                 f"below its {target:.0%} floor target")
        else:
            ok(f"Paper {P.code}: {share:.0%} of objective items are at or below A1 "
               f"({low}/{len(objective)}, target {target:.0%})")


def check_paragraph_refs():
    """"Find a word in paragraph N" must point at the paragraph the word is in,
    and the clue must not contain the answer."""
    bad = []
    for P in load_papers():
        for s in P.sections:
            for t in s.tasks:
                if not t.text:
                    continue
                paras = [x for x in t.text if x.strip()]
                for it in t.items:
                    m = re.search(r"paragraph (\d+)", it.prompt, re.I)
                    ans = it.answer.strip().strip(".").lower()
                    if m:
                        n = int(m.group(1))
                        if n < 1 or n > len(paras):
                            bad.append(f"{t.code}.{it.n}: names paragraph {n}, but the text has "
                                       f"{len(paras)}")
                        elif ans and ans not in paras[n - 1].lower():
                            where = [i + 1 for i, x in enumerate(paras) if ans in x.lower()]
                            bad.append(f"{t.code}.{it.n}: answer {it.answer!r} is not in "
                                       f"paragraph {n}" +
                                       (f" — it is in paragraph {where}" if where else ""))
                    elif re.search(r"LAST paragraph", it.prompt):
                        if ans and ans not in paras[-1].lower():
                            bad.append(f"{t.code}.{it.n}: answer {it.answer!r} is not in the "
                                       f"last paragraph")
                    # a clue must not contain its own answer
                    if re.search(r"Find a word", it.prompt, re.I) and ans and len(ans) > 3:
                        clue = re.search(r"means [“\"](.+?)[”\"]", it.prompt)
                        if clue and ans in clue.group(1).lower():
                            bad.append(f"{t.code}.{it.n}: the clue contains its own answer "
                                       f"({it.answer!r})")
    if bad:
        for b in bad:
            fail(b)
    else:
        ok("every 'find a word' item points at the right paragraph and does not give itself away")


def check_parallel():
    A, C = PAPERS["A"], PAPERS["C"]
    if len(A.sections) != len(C.sections):
        fail("Papers A and C do not have the same number of sections")
        return
    bad = []
    for a, c in zip(A.sections, C.sections):
        if (a.strand, a.marks, a.minutes) != (c.strand, c.marks, c.minutes):
            bad.append(f"{a.code}/{c.code}: {a.strand} {a.marks}/{a.minutes}′ vs "
                       f"{c.strand} {c.marks}/{c.minutes}′")
    if bad:
        fail("Papers A and C are not structurally parallel: " + "; ".join(bad))
    else:
        ok("Papers A and C are structurally parallel — same strands, marks and minutes, "
           "section by section")
    # listening speed match
    a_wpm = [int(re.search(r"\d+", DIAG_AUDIO[k].speech_rate).group()) for k in PAPER_AUDIO["A"]]
    c_wpm = [int(re.search(r"\d+", DIAG_AUDIO[k].speech_rate).group()) for k in PAPER_AUDIO["C"]]
    drift = [abs(x - y) for x, y in zip(a_wpm, c_wpm)]
    if max(drift) > 20:
        warn(f"Paper A listening speeds {a_wpm} vs Paper C {c_wpm} — drift {drift}")
    else:
        ok(f"listening speeds matched: Paper A {a_wpm} wpm vs Paper C {c_wpm} wpm")


def check_rubrics():
    for R in ALL_RUBRICS:
        if abs(R.total - R.counted) > 1e-9:
            fail(f"rubric {R.name}: stated {R.total} but criteria sum to {R.counted}")
        else:
            ok(f"rubric {R.name}: criteria sum to {R.total:g}")
        for c in R.criteria:
            if not c.descriptors:
                fail(f"rubric {R.name}, criterion {c.name}: no band descriptors")
            tops = [d[0] for d in c.descriptors]
            if str(int(c.max)) not in tops and f"{c.max:g}" not in tops:
                warn(f"rubric {R.name}, criterion {c.name}: no descriptor for the top mark "
                     f"{c.max:g}")
            if not c.vn_note:
                warn(f"rubric {R.name}, criterion {c.name}: no Vietnamese-learner note")


def check_audio():
    for key, rel in DIAG_FILES.items():
        p = os.path.join(ROOT, rel)
        if not os.path.exists(p):
            fail(f"audio {key}: {rel} is not on disk")
        elif os.path.getsize(p) < 100_000:
            fail(f"audio {key}: {rel} is only {os.path.getsize(p)} bytes")
    if not [f for f in FAILS if f.startswith("audio ")]:
        ok(f"all {len(DIAG_FILES)} diagnostic recordings are on disk")

    for key, a in DIAG_AUDIO.items():
        for fld in ("source", "licence", "attribution", "source_page", "duration", "speech_rate"):
            if not getattr(a, fld):
                fail(f"audio {key}: missing {fld}")
        if not a.audio_urls:
            fail(f"audio {key}: no source URL")
        if not a.script:
            fail(f"audio {key}: no transcript")
    if not [f for f in FAILS if "missing" in f]:
        ok("every diagnostic recording carries source, licence, credit, duration and speed")

    # every paper's audio keys exist, and every task points at a real key
    for code, keys in PAPER_AUDIO.items():
        for k in keys:
            if k not in DIAG_AUDIO:
                fail(f"Paper {code} names audio key {k}, which does not exist")
    for P in load_papers():
        for s in P.sections:
            for t in s.tasks:
                if t.audio_key and t.audio_key not in DIAG_AUDIO:
                    fail(f"{t.code}: audio_key {t.audio_key} does not exist")
    if not [f for f in FAILS if "audio key" in f or "audio_key" in f]:
        ok("every listening task points at a real recording")

    # no diagnostic recording may be reused by a teaching session
    from curriculum.audio_sources import AUDIO as COURSE_AUDIO
    course_urls = {u for a in COURSE_AUDIO.values() for u in a.audio_urls}
    clash = [k for k, a in DIAG_AUDIO.items() if set(a.audio_urls) & course_urls]
    if clash:
        fail(f"diagnostic recordings also used by teaching sessions: {clash} — a baseline "
             f"cannot use material the students are about to study")
    else:
        ok("no diagnostic recording is used by any of the 92 teaching sessions")


def check_answers_in_transcript():
    """A listening answer that is not in the recording is not an answer."""
    def norm(s):
        return re.sub(r"[^a-z0-9 ]", " ", s.lower()).split()
    misses = []
    for P in load_papers():
        for s in P.sections:
            if s.strand != "listening":
                continue
            for t in s.tasks:
                words = set(norm(" ".join(DIAG_AUDIO[t.audio_key].script)))
                for it in t.items:
                    content = [w for w in norm(it.answer)
                               if len(w) > 3 and w not in {
                                   "they", "them", "that", "this", "with", "from", "have", "what",
                                   "when", "where", "because", "which", "there", "their", "will",
                                   "would", "should", "about", "only", "into", "than", "then",
                                   "accept", "note", "both", "needed", "answer", "student",
                                   "students", "recording", "says", "said", "また"}]
                    if not content:
                        continue
                    found = sum(1 for w in content if w in words)
                    if found == 0:
                        misses.append(f"{t.code}.{it.n}: no content word of the answer occurs "
                                      f"in the transcript")
    if misses:
        for m in misses:
            warn(m)
    else:
        ok("every listening answer draws on words that really occur in its recording")


def check_adaptive():
    AD = load_adaptive()
    if AD is None:
        fail("curriculum/adaptive.py did not import")
        return
    codes = {b.code for b in AD.BRIDGES} | {e.code for e in AD.EXTENSIONS}
    lesson_codes = {l.code for l in all_lessons()}
    for t in AD.TRIGGERS:
        if not t.changes:
            fail(f"trigger {t.code}: names no changes")
        if not t.fires_when:
            fail(f"trigger {t.code}: no threshold")
        for r in t.resources:
            token = r.split(":")[0].strip()
            if re.fullmatch(r"[BE]\d", token) and token not in codes:
                fail(f"trigger {t.code} names resource {token}, which does not exist")
            if token.startswith("audio"):
                key = r.split(":", 1)[1].split("(")[0].strip()
                if key not in DIAG_AUDIO:
                    fail(f"trigger {t.code} names audio {key}, which does not exist")
        for a in t.affects:
            if re.fullmatch(r"U\d+L\d+", a) and a not in lesson_codes:
                fail(f"trigger {t.code} affects {a}, which is not a lesson in the course")
        for a in t.insert_at:
            if a not in lesson_codes:
                fail(f"trigger {t.code} inserts at {a}, which is not a lesson in the course")
            if a.startswith("D0"):
                fail(f"trigger {t.code} inserts at {a} — nothing is taught in the diagnostic "
                     f"periods, so no insert belongs there")
        if not t.insert_at and not t.standing:
            fail(f"trigger {t.code} has neither a per-lesson insert nor a standing change, so "
                 f"firing it would change nothing")
        if len(t.insert_at) > 20:
            fail(f"trigger {t.code} stamps {len(t.insert_at)} inserts — a change that belongs on "
                 f"more than twenty lesson plans is a standing change, not an insert")
    if not [f for f in FAILS if f.startswith("trigger ")]:
        ok(f"all {len(AD.TRIGGERS)} triggers have a threshold, changes, and real resources")
        ok(f"every trigger names real lessons and is either per-lesson or standing "
           f"(inserts: {sum(len(t.insert_at) for t in AD.TRIGGERS)} across all 8, "
           f"max {max(len(t.insert_at) for t in AD.TRIGGERS)} for any one)")

    # bands must partition 0..80 with no gap and no overlap
    bs = sorted(AD.BANDS, key=lambda b: b.lo)
    if bs[0].lo != 0 or bs[-1].hi != 80:
        fail(f"bands do not span 0–80: {bs[0].lo}–{bs[-1].hi}")
    for x, y in zip(bs, bs[1:]):
        if y.lo != x.hi + 1:
            fail(f"band boundary gap or overlap between {x.key} ({x.hi}) and {y.key} ({y.lo})")
    if not [f for f in FAILS if "band" in f]:
        ok("the three bands partition 0–80 with no gap and no overlap")

    # band_for must agree with the declared boundaries
    for b in AD.BANDS:
        for m in (b.lo, b.hi):
            got = AD.band_for(m).key
            if got != b.key:
                fail(f"band_for({m}) returned {got}, expected {b.key}")
    if not [f for f in FAILS if "band_for" in f]:
        ok("band_for() agrees with the declared band boundaries at every edge")

    for b in AD.BRIDGES:
        mins = sum(s.minutes for s in b.procedure)
        if mins != b.minutes:
            fail(f"bridging {b.code}: procedure totals {mins} minutes, declared {b.minutes}")
        if not b.exercises:
            fail(f"bridging {b.code}: no exercises")
        for e in b.exercises:
            if not e.answers:
                fail(f"bridging exercise {e.ref}: no answers")
        if not b.success:
            fail(f"bridging {b.code}: no exit check")
    if not [f for f in FAILS if f.startswith("bridging ")]:
        ok(f"all {len(AD.BRIDGES)} bridging lessons are 45 minutes, with exercises, answers "
           f"and an exit check")

    for e in AD.EXTENSIONS:
        if not e.steps or not e.output or not e.assess:
            fail(f"extension {e.code}: incomplete")
        if not e.demand:
            fail(f"extension {e.code}: does not say what extra demand it adds")
    if not [f for f in FAILS if f.startswith("extension ")]:
        ok(f"all {len(AD.EXTENSIONS)} extension activities name their demand, output "
           f"and assessment")

    if len(AD.CHECKPOINTS) != 4:
        warn(f"expected 4 checkpoints, found {len(AD.CHECKPOINTS)}")
    else:
        ok("4 checkpoints defined")


def check_profile():
    CP = load_profile()
    if CP is None:
        fail("curriculum/class_profile.py did not import")
        return
    if CP.DIAGNOSED:
        warn("class_profile.DIAGNOSED is True — this repository ships an adapted course. "
             "That is fine for a real class; it is not the shipped default.")
    else:
        ok("class_profile defaults to not-yet-diagnosed, so the course builds as designed")
    if CP.active_triggers():
        warn(f"undiagnosed profile still reports active triggers: {CP.active_triggers()}")
    else:
        ok("an undiagnosed profile fires no triggers and stamps no inserts")

    # the trigger engine must behave on the documented worked example
    import importlib
    m = importlib.reload(CP)
    m.DIAGNOSED = True
    m.STUDENTS = 44
    m.STRANDS = {"listening": 38, "reading": 57, "vocab": 49, "grammar": 44,
                 "writing": 41, "speaking": 46, "pron": 43}
    m.BANDS = {"foundation": 16, "core": 22, "extension": 6}
    m.SPREAD_SD = 13.8
    m.P10_P90_GAP = 34
    m.ITEM_CORRECT = {"A-G3.1": 34, "A-G3.2": 27, "A-L3.2": 31,
                      "A-G1.3": 52, "A-G1.4": 61, "A-V2.5": 45}
    m.CRITERION_MEANS = {"pron.final_consonants": 0.8, "pron.word_stress": 1.1,
                         "writing.grammar": 1.4, "writing.organisation": 0.9}
    m.WRITING_AT_OR_BELOW_4 = 34
    m.UNIT1_REVISION = 47
    fired, reasons, undecidable = m.evaluate_triggers()
    expected = ["T1", "T2", "T3", "T4", "T5", "T6", "T7"]
    if sorted(fired) != expected:
        fail(f"worked example fired {sorted(fired)}, expected {expected}")
    else:
        ok(f"the documented worked example fires exactly {', '.join(expected)} (T8 correctly "
           f"does not: 6/44 is 14%, under the 25% threshold)")
    if undecidable:
        fail(f"worked example left triggers undecidable: {undecidable}")
    if m.suggested_bridge_mode() != "warm-up":
        fail(f"worked example bridge mode {m.suggested_bridge_mode()!r}, expected 'warm-up'")
    else:
        ok("16 of 44 Foundation students correctly implies warm-up bridging delivery")

    # missing evidence must be undecidable, never a silent 'did not fire'
    m2 = importlib.reload(CP)
    m2.DIAGNOSED = True
    f2, _, u2 = m2.evaluate_triggers()
    if f2:
        fail(f"an empty profile fired {f2}")
    elif len(u2) != 8:
        fail(f"an empty profile should leave all 8 triggers undecidable, got {len(u2)}")
    else:
        ok("with no evidence entered, all 8 triggers are reported undecidable, not 'did not fire'")
    importlib.reload(CP)


def check_papers_book_has_no_answers():
    """Book 7 must not contain any answer, marking note or transcript."""
    p = os.path.join(ROOT, "output", "07_Diagnostic_Test_Papers.docx")
    if not os.path.exists(p):
        warn("output/07_Diagnostic_Test_Papers.docx not built yet — run build.py first")
        return
    try:
        from docx import Document
    except ImportError:
        warn("python-docx not available; skipped the Book 7 leak check")
        return
    d = Document(p)
    txt = "\n".join(x.text for x in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                txt += "\n" + c.text

    leaks = []
    for P in load_papers():
        for s in P.sections:
            for t in s.tasks:
                if t.audio_key:
                    for line in DIAG_AUDIO[t.audio_key].script:
                        frag = line.split(":", 1)[-1].strip()[:45]
                        if len(frag) > 25 and frag in txt:
                            leaks.append(f"transcript of {t.code}: {frag!r}")
                for it in t.items:
                    if it.note and len(it.note) > 30 and it.note[:40] in txt:
                        leaks.append(f"marking note {t.code}.{it.n}")
                    if it.tests and len(it.tests) > 12 and it.tests in txt:
                        leaks.append(f"construct label {t.code}.{it.n}")
                # an answer only counts as a leak if it is NOT quotable from the
                # task's own printed stimulus (reading answers legitimately are)
                stim = " ".join(t.text)
                for it in t.items:
                    a = (it.answer or "").strip()
                    if t.rubric or it.options or len(a) < 12:
                        continue
                    if a in txt and a not in stim:
                        leaks.append(f"answer {t.code}.{it.n}: {a[:50]!r}")
    if leaks:
        for l in leaks:
            fail("Book 7 leaks " + l)
    else:
        ok("Book 7 contains no transcript, no answer, no marking note and no construct label")


def check_outputs():
    for name in ("06_Diagnostic_and_Adaptive_System.docx", "07_Diagnostic_Test_Papers.docx"):
        p = os.path.join(ROOT, "output", name)
        if not os.path.exists(p):
            warn(f"output/{name} not built yet")
        elif os.path.getsize(p) < 20_000:
            fail(f"output/{name} is only {os.path.getsize(p)} bytes")
        else:
            ok(f"output/{name} built ({os.path.getsize(p) // 1024} KB)")


def probe_audio():
    """Re-measure every diagnostic MP3 and compare with the declared metadata."""
    for key, rel in DIAG_FILES.items():
        p = os.path.join(ROOT, rel)
        if not os.path.exists(p):
            continue
        out = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries",
             "format=duration,bit_rate:stream=codec_name,sample_rate,channels",
             "-of", "json", p], capture_output=True, text=True).stdout
        j = json.loads(out)
        dur = float(j["format"]["duration"])
        st = (j.get("streams") or [{}])[0]
        a = DIAG_AUDIO[key]
        mm, ss = (int(x) for x in a.duration.split(":"))
        declared = mm * 60 + ss
        # an excerpt's declared length is shorter than the file, by design
        if a.script_is_excerpt:
            if declared > dur + 2:
                fail(f"{key}: declares {a.duration} but the file is only {dur:.0f}s")
            else:
                ok(f"{key}: excerpt {a.duration} inside a {dur:.0f}s file · "
                   f"{st.get('codec_name')} {st.get('sample_rate')}Hz {st.get('channels')}ch")
        else:
            if abs(declared - dur) > 3:
                fail(f"{key}: declares {a.duration} ({declared}s) but the file is {dur:.0f}s")
            else:
                ok(f"{key}: {a.duration} matches the file · {st.get('codec_name')} "
                   f"{st.get('sample_rate')}Hz {st.get('channels')}ch")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--probe", action="store_true",
                    help="re-measure every MP3 with ffprobe and compare with the metadata")
    args = ap.parse_args()

    check_sequence()
    check_marks()
    check_items()
    check_calibration()
    check_paragraph_refs()
    check_parallel()
    check_rubrics()
    check_audio()
    check_answers_in_transcript()
    check_adaptive()
    check_profile()
    check_outputs()
    check_papers_book_has_no_answers()
    if args.probe:
        probe_audio()

    print("\n".join("  ✓ " + m for m in OKS))
    if WARNS:
        print("\n".join("  ! " + m for m in WARNS))
    if FAILS:
        print("\n".join("  ✗ " + m for m in FAILS))
    print(f"\n  {len(OKS)} passed · {len(WARNS)} warnings · {len(FAILS)} failures")
    sys.exit(1 if FAILS else 0)
