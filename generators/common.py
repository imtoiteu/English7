# -*- coding: utf-8 -*-
"""Shared DOCX building blocks: page setup, styles, headers, tables, boxes."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# --- palette --------------------------------------------------------------
NAVY = RGBColor(0x1F, 0x3B, 0x63)
BLUE = RGBColor(0x1E, 0x6F, 0xB8)
TEAL = RGBColor(0x0E, 0x7C, 0x66)
ORANGE = RGBColor(0xC0, 0x5A, 0x11)
GREY = RGBColor(0x59, 0x59, 0x59)
RED = RGBColor(0xB3, 0x1B, 0x1B)

FILL = {
    "objectives": "DCE9F7", "vocab": "FDF0DC", "grammar": "E4F2EC",
    "pron": "F3E8F7", "listening": "E8EEF9", "speaking": "FCE9E9",
    "reading": "EAF3E1", "writing": "FFF6DA", "note": "F2F2F2",
    "answer": "EAF6EE", "warn": "FDECEA", "head": "1F3B63",
}


# --- low-level helpers ----------------------------------------------------
def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def _borders(table, sz=4, color="9AA5B1"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def _no_borders(table):
    tbl = table._tbl
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        borders.append(el)
    tbl.tblPr.append(borders)


def field(paragraph, instr):
    """Insert a Word field (used for TOC and page numbers)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "…"
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    for el in (fldChar1, instrText, fldChar2, t, fldChar3):
        run._r.append(el)


# --- document skeleton ----------------------------------------------------
def new_doc(title, subtitle, book_name, colour=NAVY):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(1.8)

    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.08

    for name, size, col, before, after in [
        ('Heading 1', 20, colour, 14, 8), ('Heading 2', 14, BLUE, 12, 5),
        ('Heading 3', 11.5, TEAL, 9, 4), ('Heading 4', 10.5, ORANGE, 7, 3)]:
        s = doc.styles[name]
        s.font.name = 'Calibri'; s.font.size = Pt(size); s.font.bold = True
        s.font.color.rgb = col
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    # running header + footer
    hdr = sec.header.paragraphs[0]
    hdr.text = f"{title} — {book_name}"
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hdr.runs:
        r.font.size = Pt(8); r.font.color.rgb = GREY

    ftr = sec.footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ftr.add_run(f"{book_name}  |  page ")
    r.font.size = Pt(8); r.font.color.rgb = GREY
    field(ftr, "PAGE")
    for r in ftr.runs:
        r.font.size = Pt(8); r.font.color.rgb = GREY

    # --- cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(30); r.bold = True; r.font.color.rgb = colour
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = Pt(13); r.font.color.rgb = GREY
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book_name.upper()); r.font.size = Pt(20); r.bold = True; r.font.color.rgb = BLUE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("12 units  ·  94 teaching sessions  ·  four skills  ·  diagnostic-led  ·  "
                  "CEFR A1+ → A2")
    r.font.size = Pt(10.5); r.font.color.rgb = GREY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Aligned with the MOET English 7 framework (Tiếng Anh 7)")
    r.font.size = Pt(10.5); r.font.color.rgb = GREY
    doc.add_page_break()
    return doc


def add_toc(doc, levels="1-2"):
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    field(p, f'TOC \\o "{levels}" \\h \\z \\u')
    p2 = doc.add_paragraph()
    r = p2.add_run("(In Word: click in the table above and press F9 to build the contents list.)")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    doc.add_page_break()


# --- content helpers ------------------------------------------------------
def para(doc, text, bold=False, italic=False, size=10.5, color=None, align=None,
         space_after=4, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def rich(doc, parts, size=10.5, indent=0, space_after=4):
    """parts = [(text, {'b':True,'i':True,'c':RGBColor}), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, fmt in parts:
        r = p.add_run(text)
        r.bold = fmt.get('b', False); r.italic = fmt.get('i', False)
        r.font.size = Pt(fmt.get('size', size))
        if 'c' in fmt:
            r.font.color.rgb = fmt['c']
    return p


def bullets(doc, items, style="List Bullet", size=10.5, indent=0.4):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(indent + 0.3)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it); r.font.size = Pt(size)


def numbered(doc, items, size=10.5, indent=0.4):
    bullets(doc, items, style="List Number", size=size, indent=indent)


def box(doc, title, lines, fill="note", icon="", size=10, title_color=NAVY):
    """A shaded single-cell callout box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    _shade(c, FILL.get(fill, "F2F2F2"))
    _borders(t, sz=4, color="BFC9D4")
    c.paragraphs[0].text = ""
    if title:
        p = c.paragraphs[0]
        r = p.add_run(f"{icon} {title}".strip())
        r.bold = True; r.font.size = Pt(size + 0.5); r.font.color.rgb = title_color
        p.paragraph_format.space_after = Pt(2)
        first = False
    else:
        first = True
    for ln in lines:
        if first:
            p = c.paragraphs[0]; first = False
        else:
            p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(ln); r.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def table(doc, rows, header=True, widths=None, size=9.5, head_fill="1F3B63"):
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    _borders(t)
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            txt = row[j] if j < len(row) else ""
            r = p.add_run(str(txt))
            r.font.size = Pt(size)
            if header and i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cell, head_fill)
            elif i % 2 == 0:
                _shade(cell, "F7F9FB")
    if widths:
        for j, w in enumerate(widths):
            for i in range(len(rows)):
                t.cell(i, j).width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def writing_lines(doc, n=4, indent=0.4):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("." * 100)
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def rule(doc, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1F3B63')
    pbdr.append(bottom); pPr.append(pbdr)


def page_break(doc):
    doc.add_page_break()


# --- audio hyperlinks -----------------------------------------------------
def add_hyperlink(paragraph, text, target, size=9.5, color=BLUE, bold=True,
                  underline=True):
    """A real Word hyperlink (w:hyperlink + an external relationship).

    `target` is stored verbatim, so a relative path like ../audio/x.mp3 is
    resolved by Word against the document's own location on disk. That is why
    the depth passed to curriculum.audio_links.rel_path has to match where the
    file is actually written.
    """
    part = paragraph.part
    r_id = part.relate_to(target, RT.HYPERLINK, is_external=True)
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color is not None:
        c = OxmlElement('w:color'); c.set(qn('w:val'), f"{color}")
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if bold:
        rPr.append(OxmlElement('w:b'))
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


def _link_text(label, part_label, n, is_local):
    """What the clickable text says.

    A link to the local MP3 and a link to the publisher's page do different
    things, so they must not look identical — a teacher who clicks expecting
    the file and gets a web page should be able to see that coming.
    """
    if not is_local:
        base = "▶ Listen online"
        return base if n == 1 else f"{base} — {part_label}"
    return label if n == 1 else f"{label} — {part_label}"


def audio_links(doc, links, label="▶ Listen", size=9.5, indent=0.4,
                intro="", space_after=4, color=BLUE):
    """One paragraph of clickable play links.

    links = [(part_label, target, is_local), …] from audio_links.targets().
    A single-part recording gets one plain button; a multipart recording gets
    one button per part, labelled, so a teacher can play part 3 without hunting.
    """
    if not links:
        return None
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    if intro:
        r = p.add_run(intro + "  ")
        r.font.size = Pt(size); r.font.color.rgb = GREY
    for i, (part_label, target, is_local) in enumerate(links):
        if i:
            sep = p.add_run("   ")
            sep.font.size = Pt(size)
        add_hyperlink(p, _link_text(label, part_label, len(links), is_local),
                      target, size=size, color=color)
    return p


def audio_links_in_cell(cell, links, label="▶ Listen", size=9, color=BLUE):
    """Same, but inside an existing table cell (used inside callout boxes)."""
    if not links:
        return None
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    for i, (part_label, target, is_local) in enumerate(links):
        if i:
            sep = p.add_run("   ")
            sep.font.size = Pt(size)
        add_hyperlink(p, _link_text(label, part_label, len(links), is_local),
                      target, size=size, color=color)
    return p


LEVEL_NAME = {"E": "Easy", "M": "Medium", "D": "Difficult"}
LEVEL_STAR = {"E": "★☆☆", "M": "★★☆", "D": "★★★"}

# public aliases (star-import friendly)
shade = _shade
borders = _borders
no_borders = _no_borders
